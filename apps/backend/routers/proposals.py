import json

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel

from auth import _require_admin
from db import get_db_connection
from gemini import client, types
from models import ProposalGenerateRequest

router = APIRouter()


@router.post("/api/proposals/generate")
async def generate_proposal(request: ProposalGenerateRequest, _: None = Depends(_require_admin)):
    """Use Gemini to generate a complete federal proposal draft from triage data."""
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT triage_report, agency, naics, estimated_value, response_deadline, pdf_url
        FROM solicitation_queue WHERE solicitation_id=%s
    """, (request.solicitation_id,))
    sol = cur.fetchone()
    if not sol:
        cur.close()
        conn.close()
        raise HTTPException(status_code=404, detail="Solicitation not found")

    triage_report, agency, naics, est_value, deadline, pdf_url = sol

    cur.execute("""
        SELECT section, content FROM approved_language
        WHERE naics=%s OR naics IS NULL
        ORDER BY win_rate DESC NULLS LAST LIMIT 10
    """, (naics,))
    boilerplate = {r[0]: r[1] for r in cur.fetchall()}

    vendor_info = None
    if request.selected_vendor_id:
        cur.execute("""
            SELECT v.legal_name, v.cage_code, v.naics_codes, v.city, v.state,
                   v.performance_rating, v.contracts_completed, q.total_amount,
                   q.labor_rate_hourly, q.period_of_performance
            FROM vendor_registry v
            LEFT JOIN vendor_quotes q ON q.vendor_id = v.id AND q.solicitation_id=%s
            WHERE v.id=%s::uuid
        """, (request.solicitation_id, request.selected_vendor_id))
        vrow = cur.fetchone()
        if vrow:
            vendor_info = {
                "name": vrow[0], "cage": vrow[1], "naics": vrow[2],
                "location": f"{vrow[3]}, {vrow[4]}", "rating": vrow[5],
                "contracts": vrow[6], "quote_total": vrow[7],
                "labor_rate": vrow[8], "pop": vrow[9],
            }

    cur.close()
    conn.close()

    target_price = request.target_price or (float(est_value) * 1.15 if est_value else None)
    triage_summary = json.dumps(triage_report or {}, indent=2)[:2000]

    deliverable_type = "IT Services"
    section_508 = False
    if triage_report:
        scope = triage_report.get("section3_scope") or {}
        deliverable_type = scope.get("primary_deliverable_type", "IT Services")
        section_508 = (triage_report.get("section2_compliance") or {}).get("section_508_required", False)

    prompt = f"""You are the Chief Proposal Writer for Burger Consulting LLC (EIN: 84-3113166),
a federal IT services prime contractor in New York City operating under the Zero-Float doctrine.
NAICS codes: 541511 (Custom Software & Web Development), 541519 (IT Services & Project Management), 541512 (Systems Design & IT Infrastructure).
Principal: Timothy J. Burger — software engineer and federal contracting specialist with expertise in Agile delivery, Section 508 accessibility compliance, and FFP performance-based IT service contracts.

Write a complete, compelling federal proposal for the following IT services solicitation:

SOLICITATION: {request.solicitation_id}
AGENCY: {agency or 'Federal Agency'}
NAICS: {naics or 'See SOW'}
DELIVERABLE TYPE: {deliverable_type}
SECTION 508 REQUIRED: {section_508}
ESTIMATED VALUE: ${float(est_value or 0):,.0f}
RESPONSE DEADLINE: {deadline.strftime('%B %d, %Y') if deadline else 'TBD'}
TARGET PRICE: {f'${target_price:,.2f}' if target_price else 'TBD'}

TRIAGE ANALYSIS SUMMARY:
{triage_summary}

SELECTED SUBCONTRACTOR / KEY PERSONNEL: {json.dumps(vendor_info, indent=2) if vendor_info else 'TBD — to be inserted after award'}

APPROVED BOILERPLATE AVAILABLE:
{json.dumps(boilerplate, indent=2)[:1000] if boilerplate else 'None on file'}

ADDITIONAL NOTES: {request.additional_notes or 'None'}

Generate a complete proposal with these exact sections. Write 2-4 paragraphs per section using
professional federal IT proposal language (clear, direct, compliance-focused):

1. TECHNICAL_APPROACH: Describe BCG's IT delivery methodology. Include: Agile/iterative development
   approach (2-week sprints with agency checkpoints), technology stack selection rationale,
   Section 508 / WCAG 2.2 AA compliance plan (if required), security controls (NIST SP 800-53),
   remote delivery model, testing and QA protocol, and deliverable acceptance criteria.

2. MANAGEMENT_PLAN: Timothy J. Burger as Principal/PM. Communication plan (weekly status reports,
   GitHub/Jira project tracking, monthly demos to agency stakeholders). Subcontractor oversight via
   code review and deliverable QA. Risk mitigation: backup developer roster, version control,
   documentation standards. Escalation path for scope changes. Reporting cadence and tools.

3. PRICING_NARRATIVE: Labor category breakdown (Senior Developer, UI/UX Designer, PM, QA Engineer)
   at GSA IT Schedule labor rates. Hours estimate per deliverable. BCG prime management fee of
   15-20% for oversight, compliance assurance, PM, and billing administration. No upfront capital
   required — all work billed monthly or milestone-based. Competitive with market comparables
   from USASpending.gov award data for this NAICS.

4. PAST_PERFORMANCE: BCG's IT capability narrative. Note: Company in first year of federal contracting.
   Emphasize: Timothy Burger's direct software engineering expertise (UI/UX, web development, system design),
   Section 508 compliance knowledge, and the strength of the vetted subcontractor network.
   Reference relevant private-sector or academic experience. Commitment to building a strong federal track record.

5. WIN_PROBABILITY: Integer 1-100. Consider: set-aside status, competition level for this NAICS,
   BCG's technical fit for the deliverable type, Section 508 advantage (if applicable), and price competitiveness.

Return as JSON with keys: technical_approach, management_plan, pricing_narrative,
past_performance, win_probability (integer), executive_summary (2 sentences max)."""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=[prompt],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.3,
            ),
        )
        proposal_data = json.loads(response.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gemini proposal generation error: {str(e)}")

    conn = get_db_connection()
    cur = conn.cursor()
    # proposals.solicitation_id has no unique constraint, so an upsert via
    # ON CONFLICT is not available. Update the existing draft in place if one
    # exists, otherwise insert a new row — one proposal per solicitation.
    proposal_values = (
        response.text,
        proposal_data.get("technical_approach"),
        proposal_data.get("management_plan"),
        proposal_data.get("pricing_narrative"),
        proposal_data.get("past_performance"),
        proposal_data.get("win_probability"),
    )
    cur.execute("""
        UPDATE proposals SET
            gemini_draft = %s,
            technical_approach = %s,
            management_plan = %s,
            pricing_narrative = %s,
            past_performance = %s,
            win_probability = %s,
            status = 'DRAFT',
            updated_at = NOW()
        WHERE solicitation_id = %s
        RETURNING id
    """, proposal_values + (request.solicitation_id,))
    row = cur.fetchone()
    if row:
        proposal_id = row[0]
    else:
        cur.execute("""
            INSERT INTO proposals
                (solicitation_id, gemini_draft, technical_approach, management_plan,
                 pricing_narrative, past_performance, win_probability, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'DRAFT')
            RETURNING id
        """, (request.solicitation_id,) + proposal_values)
        proposal_id = cur.fetchone()[0]

    for section in boilerplate:
        cur.execute("""
            UPDATE approved_language SET times_used = times_used + 1
            WHERE naics=%s AND section=%s
        """, (naics, section))

    conn.commit()
    cur.close()
    conn.close()

    return {
        "proposal_id": str(proposal_id),
        "solicitation_id": request.solicitation_id,
        "proposal": proposal_data,
        "status": "DRAFT",
    }


@router.get("/api/proposals")
async def list_proposals(_: None = Depends(_require_admin)):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT p.id, p.solicitation_id, s.agency, s.naics, s.estimated_value,
               p.win_probability, p.status, p.created_at, p.updated_at
        FROM proposals p
        LEFT JOIN solicitation_queue s ON p.solicitation_id = s.solicitation_id
        ORDER BY p.created_at DESC
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return [{"id": str(r[0]), "solicitation_id": r[1], "agency": r[2], "naics": r[3],
             "estimated_value": float(r[4]) if r[4] else None,
             "win_probability": r[5], "status": r[6],
             "created_at": r[7].isoformat() if r[7] else None,
             "updated_at": r[8].isoformat() if r[8] else None} for r in rows]


@router.get("/api/proposals/{solicitation_id}")
async def get_proposal(solicitation_id: str, _: None = Depends(_require_admin)):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT id, solicitation_id, technical_approach, management_plan,
               pricing_narrative, past_performance, win_probability,
               status, created_at, updated_at
        FROM proposals WHERE solicitation_id=%s
        ORDER BY created_at DESC LIMIT 1
    """, (solicitation_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="No proposal found for this solicitation")
    return {"id": str(row[0]), "solicitation_id": row[1],
            "technical_approach": row[2], "management_plan": row[3],
            "pricing_narrative": row[4], "past_performance": row[5],
            "win_probability": row[6], "status": row[7],
            "created_at": row[8].isoformat() if row[8] else None,
            "updated_at": row[9].isoformat() if row[9] else None}


class _StatusUpdate(BaseModel):
    status: str


@router.put("/api/proposals/{solicitation_id}/status")
async def update_proposal_status(solicitation_id: str, request: _StatusUpdate,
                                  _: None = Depends(_require_admin)):
    status = request.status
    valid = {"DRAFT", "SUBMITTED", "AWARDED", "REJECTED"}
    if status not in valid:
        raise HTTPException(status_code=400, detail=f"status must be one of {valid}")

    # Map proposal status → solicitation pipeline phase
    sol_phase = {
        "SUBMITTED": "SUBMITTED",
        "AWARDED": "AWARDED",
        "REJECTED": "REJECTED",
    }.get(status)

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        UPDATE proposals SET status=%s, updated_at=NOW()
        WHERE solicitation_id=%s
        RETURNING id, solicitation_id
    """, (status, solicitation_id))
    row = cur.fetchone()
    if not row:
        conn.rollback()
        cur.close()
        conn.close()
        raise HTTPException(status_code=404, detail="No proposal found for this solicitation")

    if sol_phase:
        cur.execute("""
            UPDATE solicitation_queue SET phase_status=%s, updated_at=NOW()
            WHERE solicitation_id=%s
        """, (sol_phase, solicitation_id))

    conn.commit()
    cur.close()
    conn.close()
    return {"solicitation_id": solicitation_id, "status": status}


@router.get("/api/proposals/{solicitation_id}/export")
async def export_proposal_docx(solicitation_id: str, _: None = Depends(_require_admin)):
    import io
    from datetime import datetime
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT p.solicitation_id, s.agency, s.naics, s.estimated_value, s.response_deadline,
               p.technical_approach, p.management_plan, p.pricing_narrative,
               p.past_performance, p.win_probability, p.created_at
        FROM proposals p
        LEFT JOIN solicitation_queue s ON p.solicitation_id = s.solicitation_id
        WHERE p.solicitation_id = %s
        ORDER BY p.created_at DESC LIMIT 1
    """, (solicitation_id,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="No proposal found for this solicitation")

    sol_id, agency, naics, est_value, deadline, tech, mgmt, pricing, perf, win_prob, created_at = row

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    def add_heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_body(text: str):
        p = doc.add_paragraph(text or "")
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(10.5)
        return p

    title_p = doc.add_paragraph()
    title_r = title_p.add_run("BURGER CONSULTING LLC")
    title_r.bold = True
    title_r.font.size = Pt(20)
    title_r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run("Federal Procurement Proposal — Confidential")
    sub_r.font.size = Pt(10)
    sub_r.font.color.rgb = RGBColor(0x6B, 0x7A, 0x99)
    sub_p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Solicitation ID", sol_id),
        ("Issuing Agency", agency or "—"),
        ("NAICS Code", naics or "—"),
        ("Estimated Value", f"${float(est_value):,.0f}" if est_value else "—"),
        ("Response Deadline", deadline.strftime("%B %d, %Y") if deadline else "—"),
    ]
    for i, (label, value) in enumerate(meta_data):
        label_cell = meta_table.rows[i].cells[0]
        value_cell = meta_table.rows[i].cells[1]
        label_cell.text = label
        value_cell.text = value
        label_cell.paragraphs[0].runs[0].bold = True
        label_cell.paragraphs[0].runs[0].font.size = Pt(10)
        value_cell.paragraphs[0].runs[0].font.size = Pt(10)

    if win_prob is not None:
        doc.add_paragraph()
        wp = doc.add_paragraph()
        wp_r = wp.add_run(f"AI Win Probability Estimate: {win_prob}%")
        wp_r.bold = True
        wp_r.font.size = Pt(11)
        wp_r.font.color.rgb = (
            RGBColor(0x06, 0x60, 0x27) if win_prob >= 60 else RGBColor(0xD9, 0x77, 0x06)
        )

    doc.add_page_break()

    for title, content in [
        ("Technical Approach", tech),
        ("Management Plan", mgmt),
        ("Pricing Narrative", pricing),
        ("Past Performance", perf),
    ]:
        if content:
            add_heading(title, level=2)
            add_body(content)
            doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(
        f"Generated {created_at.strftime('%B %d, %Y') if created_at else 'N/A'} by Hermes AI · "
        "Burger Consulting LLC · EIN 84-3113166 · 105 E 117th St Apt 5F, New York NY 10035 · "
        "procurement@burgergov.com"
    )
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    footer_r.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"BCG_Proposal_{sol_id.replace('/', '-')}_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
